import os
import csv
import pdfplumber
from docx import Document as DocxDocument
from openpyxl import load_workbook


def get_file_type(filename):
    ext = os.path.splitext(filename)[1].lower()
    if ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff']:
        return 'image'
    elif ext == '.pdf':
        return 'pdf'
    elif ext == '.docx':
        return 'docx'
    elif ext == '.xlsx':
        return 'xlsx'
    elif ext == '.csv':
        return 'csv'
    return 'unknown'


def extract_from_pdf(file_path):
    text = ''
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + '\n'
    return text.strip()


def extract_from_docx(file_path):
    doc = DocxDocument(file_path)
    lines = []
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text.strip())
    # Also extract tables inside the docx
    for table in doc.tables:
        for row in table.rows:
            row_text = ' | '.join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                lines.append(row_text)
    return '\n'.join(lines)


def extract_from_xlsx(file_path):
    wb = load_workbook(file_path, data_only=True)
    lines = []
    for sheet in wb.worksheets:
        lines.append(f'Sheet: {sheet.title}')
        for row in sheet.iter_rows(values_only=True):
            row_text = ' | '.join(
                str(cell) for cell in row if cell is not None
            )
            if row_text.strip():
                lines.append(row_text)
    return '\n'.join(lines)


def extract_from_csv(file_path):
    lines = []
    with open(file_path, newline='', encoding='utf-8', errors='ignore') as f:
        reader = csv.reader(f)
        for row in reader:
            row_text = ' | '.join(cell.strip() for cell in row if cell.strip())
            if row_text:
                lines.append(row_text)
    return '\n'.join(lines)


def process_file(file_path, file_type):
    if file_type == 'pdf':
        return extract_from_pdf(file_path)
    elif file_type == 'docx':
        return extract_from_docx(file_path)
    elif file_type == 'xlsx':
        return extract_from_xlsx(file_path)
    elif file_type == 'csv':
        return extract_from_csv(file_path)
    return ''